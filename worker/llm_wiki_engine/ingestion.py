"""Immutable source acquisition and local document extraction."""

from __future__ import annotations

import hashlib
import json
import os
import re
import shutil
import socket
import sqlite3
import subprocess
import sys
import time
from collections.abc import Callable
from dataclasses import dataclass
from datetime import UTC, datetime
from pathlib import Path
from threading import Event

import psutil  # type: ignore[import-untyped]
import pypdfium2 as pdfium  # type: ignore[import-untyped]
from docx import Document

from .contracts import JobState

EmitEvent = Callable[[JobState, float, str, str | None, str | None, str | None], None]
SUPPORTED_SUFFIXES = {".pdf", ".docx", ".txt", ".md"}
OCR_HEARTBEAT_SECONDS = 10.0
OCR_STALL_WARNING_SECONDS = 120.0
OCR_BATCH_BASE_TIMEOUT_SECONDS = 1800.0


class IngestionCancelled(Exception):
    """Raised when the user cancels an active ingestion."""


@dataclass(frozen=True, slots=True)
class AcquiredSource:
    source_id: str
    original_name: str
    source_format: str
    content_sha256: str
    byte_size: int
    raw_path: Path


class JobLogger:
    def __init__(self, wiki_root: Path, job_id: str, emit: EmitEvent) -> None:
        self._job_id = job_id
        self._emit = emit
        log_directory = wiki_root / ".llm-wiki" / "logs"
        log_directory.mkdir(parents=True, exist_ok=True)
        self.path = log_directory / f"{job_id}.jsonl"

    def write(
        self,
        level: str,
        message: str,
        *,
        state: JobState,
        progress: float,
        source: str | None = None,
        detail: str | None = None,
    ) -> None:
        entry = {
            "timestamp": datetime.now(UTC).isoformat(),
            "level": level,
            "job_id": self._job_id,
            "state": state.value,
            "message": message,
            "source": source,
            "detail": detail,
        }
        with self.path.open("a", encoding="utf-8", newline="\n") as stream:
            stream.write(json.dumps(entry, ensure_ascii=False, separators=(",", ":")) + "\n")
        console_line = f"[LLM Wiki][{level.upper()}][{self._job_id}] {message}"
        if source:
            console_line += f" ({source})"
        if detail:
            console_line += f": {detail}"
        print(console_line, file=sys.stderr, flush=True)
        self._emit(state, progress, message, level, source, detail)


class JobProcessor:
    def __init__(
        self,
        *,
        job_id: str,
        wiki_id: str,
        wiki_root: Path,
        source_paths: list[Path],
        ocr_language: str,
        cancellation: Event,
        emit: EmitEvent,
    ) -> None:
        self._job_id = job_id
        self._wiki_id = wiki_id
        self._wiki_root = wiki_root.resolve()
        self._source_paths = source_paths
        self._ocr_language = normalize_ocr_languages(ocr_language)
        self._cancellation = cancellation
        self._emit = emit
        self._logger = JobLogger(self._wiki_root, job_id, emit)

    def run(self) -> int:
        self._logger.write(
            "info",
            "job.started",
            state=JobState.ACQUIRING,
            progress=0.0,
            detail=f"sources={len(self._source_paths)}",
        )
        acquired = [
            self._acquire(path, index, len(self._source_paths))
            for index, path in enumerate(self._source_paths, start=1)
        ]
        self._check_cancelled()

        pdf_sources = [source for source in acquired if source.source_format == "pdf"]
        direct_sources = [source for source in acquired if source.source_format != "pdf"]
        processed = 0
        for index, source in enumerate(direct_sources, start=1):
            progress = 0.35 + (0.35 * index / max(len(acquired), 1))
            self._extract_direct(source, progress)
            processed += 1
        if pdf_sources:
            digital_pdfs: list[AcquiredSource] = []
            scanned_pdfs: list[AcquiredSource] = []
            for source in pdf_sources:
                destination = (
                    digital_pdfs if pdf_has_embedded_text(source.raw_path) else scanned_pdfs
                )
                destination.append(source)
                self._logger.write(
                    "info",
                    "pdf.digital_detected" if destination is digital_pdfs else "pdf.ocr_required",
                    state=JobState.EXTRACTING,
                    progress=0.7,
                    source=source.original_name,
                )
            if digital_pdfs:
                self._extract_digital_pdfs(digital_pdfs, 0.72)
            if scanned_pdfs:
                self._extract_ocr_pdfs(scanned_pdfs, 0.76)
            processed += len(pdf_sources)

        self._validate(acquired)
        self._logger.write(
            "info",
            "job.completed",
            state=JobState.COMPLETED,
            progress=1.0,
            detail=f"processed={processed}",
        )
        return processed

    def log_failure(self, detail: str) -> None:
        self._logger.write(
            "error",
            "job.failed",
            state=JobState.FAILED,
            progress=0.0,
            detail=detail,
        )

    def log_cancelled(self) -> None:
        self._logger.write(
            "warning",
            "job.cancelled",
            state=JobState.CANCELLED,
            progress=0.0,
        )

    def _acquire(self, source_path: Path, index: int, total: int) -> AcquiredSource:
        self._check_cancelled()
        resolved = source_path.resolve(strict=True)
        suffix = resolved.suffix.lower()
        if not resolved.is_file() or suffix not in SUPPORTED_SUFFIXES:
            raise ValueError(f"Unsupported or missing source: {resolved.name}")
        byte_size = resolved.stat().st_size
        if byte_size > 2 * 1024 * 1024 * 1024:
            raise ValueError(f"Source exceeds the 2 GB limit: {resolved.name}")
        content_hash = sha256_file(resolved, self._cancellation)
        source_id = content_hash
        raw_directory = self._wiki_root / ".llm-wiki" / "raw" / content_hash
        raw_directory.mkdir(parents=True, exist_ok=True)
        safe_name = sanitize_filename(resolved.name)
        raw_path = raw_directory / safe_name
        if not raw_path.exists():
            temporary = raw_directory / f".{safe_name}.{self._job_id}.tmp"
            shutil.copyfile(resolved, temporary)
            os.replace(temporary, raw_path)

        acquired = AcquiredSource(
            source_id=source_id,
            original_name=resolved.name,
            source_format=suffix.removeprefix("."),
            content_sha256=content_hash,
            byte_size=byte_size,
            raw_path=raw_path,
        )
        self._record_source(acquired)
        progress = 0.3 * index / total
        self._logger.write(
            "info",
            "source.acquired",
            state=JobState.ACQUIRING,
            progress=progress,
            source=resolved.name,
            detail=f"sha256={content_hash[:12]} bytes={byte_size}",
        )
        return acquired

    def _extract_direct(self, source: AcquiredSource, progress: float) -> None:
        self._check_cancelled()
        if source.source_format == "docx":
            markdown = docx_to_markdown(source.raw_path)
            extractor = "python-docx"
        else:
            markdown = source.raw_path.read_text(encoding="utf-8-sig", errors="replace")
            extractor = "direct-text"
        if not markdown.strip():
            raise ValueError(f"No readable content extracted from {source.original_name}")
        self._write_artifacts(source, markdown, extractor)
        self._logger.write(
            "info",
            "source.extracted",
            state=JobState.EXTRACTING,
            progress=progress,
            source=source.original_name,
            detail=f"extractor={extractor}",
        )

    def _extract_digital_pdfs(self, sources: list[AcquiredSource], progress: float) -> None:
        """Extract selectable PDF content with OpenDataLoader's fast structural parser."""
        self._check_cancelled()
        job_directory = self._wiki_root / ".llm-wiki" / "artifacts" / self._job_id
        staged_sources = stage_pdf_sources(sources, job_directory / "digital-pdf-input")
        output_directory = job_directory / "digital-pdf-output"
        output_directory.mkdir(parents=True, exist_ok=True)
        client_executable = Path(sys.executable).with_name("opendataloader-pdf.exe")
        if not client_executable.is_file():
            raise RuntimeError("OpenDataLoader PDF executable is not installed")
        page_counts = [pdf_page_count(path) for path in staged_sources]
        total_pages = sum(page_counts)
        self._logger.write(
            "info",
            "pdf.direct_batch_started",
            state=JobState.EXTRACTING,
            progress=progress,
            detail=f"documents={len(sources)} pages={total_pages} ocr=false single_jvm=true",
        )
        log_path = self._logger.path.with_name(f"{self._job_id}-pdf-parser.log")
        with log_path.open("ab") as process_log:
            process = subprocess.Popen(
                build_pdf_direct_batch_command(client_executable, staged_sources, output_directory),
                stdin=subprocess.DEVNULL,
                stdout=process_log,
                stderr=subprocess.STDOUT,
                creationflags=hidden_process_flags(),
            )
            monitor_conversion_process(process, self._cancellation, total_pages)
        if process.returncode != 0:
            raise RuntimeError(
                f"OpenDataLoader exited with code {process.returncode}; see {log_path.name}"
            )
        for index, (source, staged_path) in enumerate(
            zip(sources, staged_sources, strict=True), start=1
        ):
            markdown, semantic_json = read_pdf_outputs(output_directory, staged_path, source)
            self._write_artifacts(
                source,
                markdown,
                "opendataloader-pdf-structured-text",
                semantic_json=semantic_json,
            )
            self._logger.write(
                "info",
                "source.text_extracted",
                state=JobState.EXTRACTING,
                progress=progress + (0.04 * index / len(sources)),
                source=source.original_name,
                detail="ocr=false format=markdown,json",
            )

    def _extract_ocr_pdfs(self, sources: list[AcquiredSource], progress: float) -> None:
        self._check_cancelled()
        job_directory = self._wiki_root / ".llm-wiki" / "artifacts" / self._job_id
        input_directory = job_directory / "pdf-input"
        output_directory = job_directory / "pdf-output"
        output_directory.mkdir(parents=True, exist_ok=True)
        staged_sources = stage_pdf_sources(sources, input_directory)
        page_counts = [pdf_page_count(path) for path in staged_sources]
        total_pages = sum(page_counts)
        if total_pages == 0:
            raise ValueError("The selected PDFs contain no pages")
        port = available_local_port()
        server_log_path = self._logger.path.with_name(f"{self._job_id}-ocr-server.log")
        server_executable = Path(sys.executable).with_name("opendataloader-pdf-hybrid.exe")
        client_executable = Path(sys.executable).with_name("opendataloader-pdf.exe")
        if not server_executable.is_file() or not client_executable.is_file():
            raise RuntimeError("OpenDataLoader hybrid executables are not installed")

        ocr_device = preferred_ocr_device()
        if ocr_device == "cpu" and nvidia_gpu_is_present():
            self._logger.write(
                "warning",
                "ocr.gpu_runtime_missing",
                state=JobState.EXTRACTING,
                progress=progress,
                detail="nvidia_gpu=true torch_cuda=false fallback=cpu",
            )

        self._logger.write(
            "info",
            "ocr.server_starting",
            state=JobState.EXTRACTING,
            progress=progress,
            detail=(f"force_ocr=true languages={self._ocr_language} device={ocr_device}"),
        )
        with server_log_path.open("ab") as server_log:
            server = subprocess.Popen(
                [
                    str(server_executable),
                    "--host",
                    "127.0.0.1",
                    "--port",
                    str(port),
                    "--force-ocr",
                    "--ocr-lang",
                    self._ocr_language,
                    "--device",
                    ocr_device,
                ],
                stdin=subprocess.DEVNULL,
                stdout=server_log,
                stderr=subprocess.STDOUT,
                creationflags=hidden_process_flags(),
            )
            try:
                wait_for_server(server, port, self._cancellation, timeout_seconds=300)
                self._logger.write(
                    "info",
                    "ocr.server_ready",
                    state=JobState.EXTRACTING,
                    progress=progress,
                )
                self._logger.write(
                    "info",
                    "ocr.batch_started",
                    state=JobState.EXTRACTING,
                    progress=progress,
                    detail=f"documents={len(sources)} pages={total_pages} single_jvm=true",
                )
                log_offset = server_log_path.stat().st_size
                command = build_pdf_batch_command(
                    client_executable, staged_sources, output_directory, port
                )
                client = subprocess.Popen(
                    command,
                    stdin=subprocess.DEVNULL,
                    stdout=server_log,
                    stderr=subprocess.STDOUT,
                    creationflags=hidden_process_flags(),
                )
                monitor_ocr_process(
                    client=client,
                    server=server,
                    cancellation=self._cancellation,
                    log_path=server_log_path,
                    log_offset=log_offset,
                    logger=self._logger,
                    progress=progress,
                    sources=sources,
                    page_counts=page_counts,
                )
                if client.returncode != 0:
                    raise RuntimeError(
                        "OpenDataLoader exited with code "
                        f"{client.returncode}; see {server_log_path.name}"
                    )

                completed_pages = 0
                for document_index, (source, staged_path, page_count) in enumerate(
                    zip(sources, staged_sources, page_counts, strict=True), start=1
                ):
                    markdown, semantic_json = read_pdf_outputs(
                        output_directory, staged_path, source
                    )
                    self._write_artifacts(
                        source,
                        markdown,
                        "opendataloader-pdf-hybrid-force-ocr",
                        semantic_json=semantic_json,
                    )
                    completed_pages += page_count
                    self._logger.write(
                        "info",
                        "source.ocr_completed",
                        state=JobState.EXTRACTING,
                        progress=ocr_progress(completed_pages, total_pages),
                        source=source.original_name,
                        detail=f"document={document_index}/{len(sources)} pages={page_count}",
                    )
            finally:
                terminate_process(server)

    def _write_artifacts(
        self,
        source: AcquiredSource,
        markdown: str,
        extractor: str,
        *,
        semantic_json: object | None = None,
    ) -> None:
        artifact_directory = self._wiki_root / ".llm-wiki" / "artifacts" / source.content_sha256
        artifact_directory.mkdir(parents=True, exist_ok=True)
        atomic_write_text(artifact_directory / "document.md", markdown)
        if semantic_json is not None:
            atomic_write_text(
                artifact_directory / "semantic.json",
                json.dumps(semantic_json, ensure_ascii=False, indent=2) + "\n",
            )
        manifest = {
            "schema_version": "1.0",
            "source_id": source.source_id,
            "original_name": source.original_name,
            "source_format": source.source_format,
            "content_sha256": source.content_sha256,
            "byte_size": source.byte_size,
            "extractor": extractor,
            "semantic_json": semantic_json is not None,
        }
        atomic_write_text(
            artifact_directory / "manifest.json",
            json.dumps(manifest, ensure_ascii=False, indent=2) + "\n",
        )
        note_name = f"{slugify(Path(source.original_name).stem)}-{source.content_sha256[:8]}.md"
        note = (
            "---\n"
            f"source_id: {source.source_id}\n"
            f"original_name: {json.dumps(source.original_name, ensure_ascii=False)}\n"
            f"source_format: {source.source_format}\n"
            f"extractor: {extractor}\n"
            "---\n\n"
            f"{markdown.strip()}\n"
        )
        atomic_write_text(self._wiki_root / "sources" / note_name, note)

    def _validate(self, acquired: list[AcquiredSource]) -> None:
        self._logger.write(
            "info",
            "job.validating",
            state=JobState.VALIDATING,
            progress=0.94,
        )
        missing = [
            source.original_name
            for source in acquired
            if not (
                self._wiki_root / ".llm-wiki" / "artifacts" / source.content_sha256 / "document.md"
            ).is_file()
        ]
        if missing:
            raise RuntimeError(f"Missing extraction artifacts: {', '.join(missing)}")

    def _record_source(self, source: AcquiredSource) -> None:
        database_path = self._wiki_root / ".llm-wiki" / "catalog.sqlite3"
        with sqlite3.connect(database_path, timeout=30) as connection:
            connection.execute("PRAGMA foreign_keys=ON")
            connection.execute(
                "INSERT OR REPLACE INTO source_records "
                "(source_id, job_id, original_name, source_format, content_sha256, byte_size) "
                "VALUES (?, ?, ?, ?, ?, ?)",
                (
                    source.source_id,
                    self._job_id,
                    source.original_name,
                    source.source_format,
                    source.content_sha256,
                    source.byte_size,
                ),
            )

    def _check_cancelled(self) -> None:
        if self._cancellation.is_set():
            raise IngestionCancelled


def sha256_file(path: Path, cancellation: Event) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        while chunk := stream.read(1024 * 1024):
            if cancellation.is_set():
                raise IngestionCancelled
            digest.update(chunk)
    return digest.hexdigest()


def docx_to_markdown(path: Path) -> str:
    document = Document(str(path))
    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style = paragraph.style.name.lower() if paragraph.style is not None else ""
        if style.startswith("heading"):
            digits = "".join(character for character in style if character.isdigit())
            level = min(max(int(digits or "1"), 1), 6)
            lines.append(f"{'#' * level} {text}")
        elif "list" in style:
            lines.append(f"- {text}")
        else:
            lines.append(text)
    for table in document.tables:
        rows = [[cell.text.strip().replace("\n", " ") for cell in row.cells] for row in table.rows]
        if not rows:
            continue
        lines.append("| " + " | ".join(rows[0]) + " |")
        lines.append("| " + " | ".join("---" for _ in rows[0]) + " |")
        lines.extend("| " + " | ".join(row) + " |" for row in rows[1:])
    return "\n\n".join(lines).strip() + "\n"


def available_local_port() -> int:
    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as listener:
        listener.bind(("127.0.0.1", 0))
        return int(listener.getsockname()[1])


def pdf_page_count(path: Path) -> int:
    document = pdfium.PdfDocument(str(path))
    try:
        return len(document)
    finally:
        document.close()


def pdf_has_embedded_text(path: Path) -> bool:
    """Return true as soon as a PDF exposes selectable alphanumeric text."""
    document = pdfium.PdfDocument(str(path))
    try:
        for page in document:
            text_page = page.get_textpage()
            try:
                if any(character.isalnum() for character in text_page.get_text_range()):
                    return True
            finally:
                text_page.close()
                page.close()
        return False
    finally:
        document.close()


def stage_pdf_sources(sources: list[AcquiredSource], input_directory: Path) -> list[Path]:
    input_directory.mkdir(parents=True, exist_ok=True)
    staged_sources: list[Path] = []
    for index, source in enumerate(sources, start=1):
        staged_path = input_directory / (
            f"{index:04d}-{source.content_sha256[:12]}-{sanitize_filename(source.original_name)}"
        )
        if not staged_path.exists():
            try:
                os.link(source.raw_path, staged_path)
            except OSError:
                shutil.copyfile(source.raw_path, staged_path)
        staged_sources.append(staged_path)
    return staged_sources


def read_pdf_outputs(
    output_directory: Path, staged_path: Path, source: AcquiredSource
) -> tuple[str, object | None]:
    markdown_path = find_pdf_output(output_directory, staged_path.stem, ".md")
    if markdown_path is None:
        raise RuntimeError(f"OpenDataLoader did not produce Markdown for {source.original_name}")
    markdown = markdown_path.read_text(encoding="utf-8", errors="replace").strip()
    if not markdown:
        raise RuntimeError(f"OpenDataLoader produced empty Markdown for {source.original_name}")
    semantic_path = find_pdf_output(output_directory, staged_path.stem, ".json")
    semantic_json = (
        json.loads(semantic_path.read_text(encoding="utf-8")) if semantic_path is not None else None
    )
    return markdown, semantic_json


def preferred_ocr_device() -> str:
    try:
        import torch

        return "cuda" if torch.cuda.is_available() else "cpu"
    except (ImportError, OSError):
        return "cpu"


def nvidia_gpu_is_present() -> bool:
    executable = shutil.which("nvidia-smi")
    if executable is None:
        return False
    try:
        result = subprocess.run(
            [executable, "--query-gpu=name", "--format=csv,noheader"],
            stdin=subprocess.DEVNULL,
            stdout=subprocess.PIPE,
            stderr=subprocess.DEVNULL,
            timeout=5,
            check=False,
            creationflags=hidden_process_flags(),
        )
    except (OSError, subprocess.TimeoutExpired):
        return False
    return result.returncode == 0 and bool(result.stdout.strip())


def ocr_progress(completed_pages: int, total_pages: int) -> float:
    return 0.72 + (0.16 * completed_pages / max(total_pages, 1))


def build_pdf_batch_command(
    client_executable: Path,
    staged_sources: list[Path],
    output_directory: Path,
    port: int,
) -> list[str]:
    return [
        str(client_executable),
        *[str(path) for path in staged_sources],
        "--output-dir",
        str(output_directory),
        "--format",
        "markdown,json",
        "--markdown-page-separator",
        "<!-- page: %page-number% -->",
        "--hybrid",
        "docling-fast",
        "--hybrid-url",
        f"http://127.0.0.1:{port}",
        "--hybrid-mode",
        "full",
    ]


def build_pdf_direct_batch_command(
    client_executable: Path,
    staged_sources: list[Path],
    output_directory: Path,
) -> list[str]:
    return [
        str(client_executable),
        *[str(path) for path in staged_sources],
        "--output-dir",
        str(output_directory),
        "--format",
        "markdown,json",
        "--markdown-page-separator",
        "<!-- page: %page-number% -->",
    ]


def monitor_conversion_process(
    process: subprocess.Popen[bytes], cancellation: Event, total_pages: int
) -> None:
    started_at = time.monotonic()
    timeout_seconds = max(120.0, total_pages * 10.0)
    while process.poll() is None:
        if cancellation.wait(0.1):
            terminate_process(process)
            raise IngestionCancelled
        if time.monotonic() - started_at >= timeout_seconds:
            terminate_process(process)
            raise TimeoutError("Structured PDF extraction exceeded its safety timeout")


@dataclass(frozen=True, slots=True)
class OcrMetrics:
    cpu_percent: float
    memory_mb: float


class ProcessTreeSampler:
    def __init__(self, root_process_ids: list[int]) -> None:
        self._root_process_ids = root_process_ids
        self._previous_cpu: dict[int, float] = {}
        self._previous_time = time.monotonic()

    def sample(self) -> OcrMetrics:
        processes: dict[int, psutil.Process] = {}
        for process_id in self._root_process_ids:
            try:
                root = psutil.Process(process_id)
                processes[root.pid] = root
                processes.update({child.pid: child for child in root.children(recursive=True)})
            except (psutil.NoSuchProcess, psutil.AccessDenied):
                continue
        current_cpu: dict[int, float] = {}
        memory_bytes = 0
        for process_id, process in processes.items():
            try:
                times = process.cpu_times()
                current_cpu[process_id] = float(times.user + times.system)
                memory_bytes += process.memory_info().rss
            except (psutil.NoSuchProcess, psutil.AccessDenied):
                continue
        now = time.monotonic()
        elapsed = max(now - self._previous_time, 0.001)
        cpu_delta = sum(
            max(value - self._previous_cpu.get(process_id, value), 0.0)
            for process_id, value in current_cpu.items()
        )
        logical_cpus = psutil.cpu_count() or 1
        cpu_percent = min((cpu_delta / elapsed / logical_cpus) * 100.0, 100.0)
        self._previous_cpu = current_cpu
        self._previous_time = now
        return OcrMetrics(cpu_percent=cpu_percent, memory_mb=memory_bytes / 1024 / 1024)


def monitor_ocr_process(
    *,
    client: subprocess.Popen[bytes],
    server: subprocess.Popen[bytes],
    cancellation: Event,
    log_path: Path,
    log_offset: int,
    logger: JobLogger,
    progress: float,
    sources: list[AcquiredSource],
    page_counts: list[int],
) -> None:
    started_at = time.monotonic()
    last_activity_at = started_at
    last_heartbeat_at = started_at
    last_warning_at = 0.0
    sampler = ProcessTreeSampler([client.pid, server.pid])
    sampler.sample()
    current_offset = log_offset
    current_document = 0
    total_pages = sum(page_counts)

    while client.poll() is None:
        if cancellation.wait(0.25):
            terminate_process(client)
            raise IngestionCancelled
        current_offset, entries = read_ocr_log_entries(log_path, current_offset)
        now = time.monotonic()
        if entries:
            last_activity_at = now
        for level, message, detail in entries:
            source = sources[min(current_document, len(sources) - 1)].original_name
            event_progress = ocr_progress(sum(page_counts[:current_document]), total_pages)
            if message == "ocr.backend_document_completed":
                event_progress = ocr_progress(sum(page_counts[: current_document + 1]), total_pages)
            logger.write(
                level,
                message,
                state=JobState.EXTRACTING,
                progress=max(progress, event_progress),
                source=source,
                detail=detail,
            )
            if message == "ocr.backend_document_completed":
                current_document = min(current_document + 1, len(sources) - 1)
        if now - last_heartbeat_at >= OCR_HEARTBEAT_SECONDS:
            metrics = sampler.sample()
            if metrics.cpu_percent >= 0.5:
                last_activity_at = now
            elapsed_seconds = int(now - started_at)
            source = sources[min(current_document, len(sources) - 1)].original_name
            completed_pages = sum(page_counts[:current_document])
            logger.write(
                "info",
                "ocr.working",
                state=JobState.EXTRACTING,
                progress=max(progress, ocr_progress(completed_pages, total_pages)),
                source=source,
                detail=(
                    f"document={current_document + 1}/{len(sources)} "
                    f"pages={page_counts[current_document]} "
                    f"elapsed={format_duration(elapsed_seconds)} "
                    f"cpu={metrics.cpu_percent:.1f}% memory={metrics.memory_mb:.0f}MB"
                ),
            )
            last_heartbeat_at = now
        inactive_seconds = now - last_activity_at
        if (
            inactive_seconds >= OCR_STALL_WARNING_SECONDS
            and now - last_warning_at >= OCR_STALL_WARNING_SECONDS
        ):
            logger.write(
                "warning",
                "ocr.possible_stall",
                state=JobState.EXTRACTING,
                progress=max(
                    progress,
                    ocr_progress(sum(page_counts[:current_document]), total_pages),
                ),
                source=sources[current_document].original_name,
                detail=f"no_activity={format_duration(int(inactive_seconds))}",
            )
            last_warning_at = now
        timeout_seconds = max(OCR_BATCH_BASE_TIMEOUT_SECONDS, total_pages * 120.0)
        if now - started_at >= timeout_seconds:
            terminate_process(client)
            raise TimeoutError(
                "OCR batch exceeded its safety timeout "
                f"({len(sources)} documents, {total_pages} pages)"
            )

    _, entries = read_ocr_log_entries(log_path, current_offset)
    for level, message, detail in entries:
        source = sources[min(current_document, len(sources) - 1)].original_name
        event_progress = ocr_progress(sum(page_counts[:current_document]), total_pages)
        if message == "ocr.backend_document_completed":
            event_progress = ocr_progress(sum(page_counts[: current_document + 1]), total_pages)
        logger.write(
            level,
            message,
            state=JobState.EXTRACTING,
            progress=max(progress, event_progress),
            source=source,
            detail=detail,
        )
        if message == "ocr.backend_document_completed":
            current_document = min(current_document + 1, len(sources) - 1)


def read_ocr_log_entries(path: Path, offset: int) -> tuple[int, list[tuple[str, str, str]]]:
    if not path.is_file():
        return offset, []
    with path.open("rb") as stream:
        stream.seek(offset)
        content = stream.read()
    last_line_break = content.rfind(b"\n")
    if last_line_break < 0:
        return offset, []
    complete = content[: last_line_break + 1]
    new_offset = offset + len(complete)
    entries = [
        entry
        for line in complete.decode("utf-8", errors="replace").splitlines()
        if (entry := classify_ocr_log_line(line)) is not None
    ]
    return new_offset, entries


def classify_ocr_log_line(line: str) -> tuple[str, str, str] | None:
    compact = " ".join(line.strip().split())
    if not compact:
        return None
    lowered = compact.lower()
    detail = compact[-500:]
    if "downloading" in lowered and "model" in lowered:
        return "info", "ocr.models_downloading", detail
    if "download complete" in lowered:
        return "info", "ocr.model_downloaded", detail
    if "engine ready" in lowered or "pipeline initialized" in lowered:
        return "info", "ocr.model_ready", detail
    if "accelerator device" in lowered:
        return "info", "ocr.accelerator", detail
    if re.search(r"processing document\b", lowered):
        return "info", "ocr.backend_processing", detail
    if "finished converting document" in lowered:
        return "info", "ocr.backend_document_completed", detail
    if "starting hybrid processing for" in lowered or "pages via docling-fast" in lowered:
        return "info", "ocr.backend_pages", detail
    if "warning" in lowered and "pin_memory" not in lowered:
        return "warning", "ocr.backend_warning", detail
    if "error" in lowered or "traceback" in lowered:
        return "error", "ocr.backend_error", detail
    return None


def format_duration(seconds: int) -> str:
    minutes, remaining_seconds = divmod(max(seconds, 0), 60)
    return f"{minutes:02d}:{remaining_seconds:02d}"


def wait_for_server(
    process: subprocess.Popen[bytes], port: int, cancellation: Event, *, timeout_seconds: int
) -> None:
    deadline = time.monotonic() + timeout_seconds
    while time.monotonic() < deadline:
        if cancellation.is_set():
            terminate_process(process)
            raise IngestionCancelled
        if process.poll() is not None:
            raise RuntimeError(f"OpenDataLoader OCR server exited with code {process.returncode}")
        with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as probe:
            probe.settimeout(0.25)
            if probe.connect_ex(("127.0.0.1", port)) == 0:
                return
        time.sleep(0.25)
    raise TimeoutError("OpenDataLoader OCR server did not become ready within 5 minutes")


def terminate_process(process: subprocess.Popen[bytes]) -> None:
    if process.poll() is not None:
        return
    process.terminate()
    try:
        process.wait(timeout=5)
    except subprocess.TimeoutExpired:
        process.kill()
        process.wait(timeout=5)


def find_pdf_output(directory: Path, source_stem: str, suffix: str) -> Path | None:
    candidates = [path for path in directory.rglob(f"*{suffix}") if path.stem == source_stem]
    if not candidates:
        candidates = list(directory.rglob(f"{source_stem}*{suffix}"))
    return candidates[0] if candidates else None


def normalize_ocr_languages(value: str) -> str:
    mapping = {"ita": "it", "eng": "en"}
    values = value.replace("+", ",").split(",")
    normalized = [mapping.get(item.strip().lower(), item.strip().lower()) for item in values]
    return ",".join(item for item in normalized if item) or "it,en"


def sanitize_filename(value: str) -> str:
    forbidden = '<>:"/\\|?*'
    cleaned = "".join("_" if character in forbidden else character for character in value)
    return cleaned.strip(" .")[:180] or "document"


def slugify(value: str) -> str:
    cleaned = "".join(character.lower() if character.isalnum() else "-" for character in value)
    return "-".join(part for part in cleaned.split("-") if part)[:80] or "document"


def atomic_write_text(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_name(f".{path.name}.{os.getpid()}.tmp")
    temporary.write_text(content, encoding="utf-8", newline="\n")
    os.replace(temporary, path)


def hidden_process_flags() -> int:
    return int(getattr(subprocess, "CREATE_NO_WINDOW", 0))
