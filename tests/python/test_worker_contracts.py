from __future__ import annotations

import json
import sqlite3
from io import StringIO
from pathlib import Path
from typing import Any, cast

from docx import Document
from jsonschema import Draft202012Validator, FormatChecker
from llm_wiki_engine.cli import run
from llm_wiki_engine.contracts import (
    ERROR_CATEGORIES,
    INGEST_STRATEGIES,
    JOB_STATES,
    MESSAGE_TYPES,
    PROVIDER_IDS,
    REVIEW_SEVERITIES,
    SOURCE_FORMATS,
    IpcEnvelope,
    ProviderOperationState,
    ProviderStatus,
    ProviderTransport,
)
from llm_wiki_engine.ingestion import (
    build_pdf_batch_command,
    build_pdf_direct_batch_command,
    classify_ocr_log_line,
    clean_extracted_text,
    ocr_progress,
    read_ocr_log_entries,
)

REPOSITORY_ROOT = Path(__file__).parents[2]
FIXTURES = REPOSITORY_ROOT / "tests" / "fixtures" / "contracts"
SCHEMAS = REPOSITORY_ROOT / "schemas" / "v1"


def read_json(name: str) -> dict[str, Any]:
    return cast(dict[str, Any], json.loads((FIXTURES / name).read_text(encoding="utf-8")))


def test_python_enums_match_language_neutral_fixture() -> None:
    values = read_json("contract-enums.json")

    assert values["message_types"] == list(MESSAGE_TYPES)
    assert values["job_states"] == list(JOB_STATES)
    assert values["source_formats"] == list(SOURCE_FORMATS)
    assert values["error_categories"] == list(ERROR_CATEGORIES)
    assert values["review_severities"] == list(REVIEW_SEVERITIES)
    assert values["provider_ids"] == list(PROVIDER_IDS)
    assert values["provider_transports"] == list(ProviderTransport)
    assert values["provider_statuses"] == list(ProviderStatus)
    assert values["provider_operation_states"] == list(ProviderOperationState)
    assert values["ingest_strategies"] == list(INGEST_STRATEGIES)


def test_all_shared_examples_validate_against_their_schema() -> None:
    examples = cast(
        list[dict[str, str]],
        json.loads((FIXTURES / "schema-examples.json").read_text(encoding="utf-8")),
    )

    for entry in examples:
        schema = read_schema(entry["schema"])
        Draft202012Validator.check_schema(schema)
        Draft202012Validator(schema, format_checker=FormatChecker()).validate(
            read_json(entry["example"])
        )


def read_schema(name: str) -> dict[str, Any]:
    return cast(dict[str, Any], json.loads((SCHEMAS / name).read_text(encoding="utf-8")))


def test_parses_shared_ipc_fixture() -> None:
    envelope = IpcEnvelope.from_mapping(read_json("ipc-request.json"))

    assert envelope.protocol_version == "1.0"
    assert envelope.payload == {"action": "health"}


def test_worker_returns_a_versioned_ready_response() -> None:
    input_stream = StringIO(json.dumps(read_json("ipc-request.json")) + "\n")
    output_stream = StringIO()

    assert run(input_stream, output_stream) == 0
    response = json.loads(output_stream.getvalue())
    Draft202012Validator(read_schema("ipc-envelope.schema.json")).validate(response)
    assert response["protocol_version"] == "1.0"
    assert response["message_type"] == "response"
    assert response["payload"]["status"] == "ready"


def test_worker_streams_progress_and_completion_for_a_job() -> None:
    request = {
        "protocol_version": "1.0",
        "message_type": "request",
        "request_id": "request-1",
        "wiki_id": "wiki-1",
        "job_id": "job-1",
        "payload": {"action": "start_job", "steps": 3, "delay_ms": 1, "source_count": 2},
    }
    output_stream = StringIO()

    assert run(StringIO(json.dumps(request) + "\n"), output_stream) == 0
    responses = [json.loads(line) for line in output_stream.getvalue().splitlines()]

    assert [response["message_type"] for response in responses] == [
        "progress",
        "progress",
        "progress",
        "response",
    ]
    assert responses[-1]["payload"] == {"status": "completed", "processed_sources": 2}


def test_worker_can_cancel_an_active_job() -> None:
    start = {
        "protocol_version": "1.0",
        "message_type": "request",
        "request_id": "request-start",
        "wiki_id": "wiki-1",
        "job_id": "job-1",
        "payload": {"action": "start_job", "steps": 20, "delay_ms": 100},
    }
    cancel = {
        "protocol_version": "1.0",
        "message_type": "request",
        "request_id": "request-cancel",
        "wiki_id": "wiki-1",
        "payload": {"action": "cancel_job", "job_id": "job-1"},
    }
    output_stream = StringIO()

    assert run(StringIO(json.dumps(start) + "\n" + json.dumps(cancel) + "\n"), output_stream) == 0
    responses = [json.loads(line) for line in output_stream.getvalue().splitlines()]

    assert any(
        response["request_id"] == "request-cancel"
        and response["payload"]["status"] == "cancellation_requested"
        for response in responses
    )
    assert any(
        response["request_id"] == "request-start"
        and response["message_type"] == "error"
        and response["payload"]["category"] == "cancelled"
        for response in responses
    )


def test_worker_acquires_and_extracts_real_text_and_markdown(tmp_path: Path) -> None:
    wiki_root = tmp_path / "wiki"
    (wiki_root / ".llm-wiki").mkdir(parents=True)
    (wiki_root / "sources").mkdir()
    database = sqlite3.connect(wiki_root / ".llm-wiki" / "catalog.sqlite3")
    database.executescript(
        "CREATE TABLE jobs (job_id TEXT PRIMARY KEY);"
        "CREATE TABLE source_records ("
        "source_id TEXT PRIMARY KEY, job_id TEXT NOT NULL REFERENCES jobs(job_id), "
        "original_name TEXT NOT NULL, source_format TEXT NOT NULL, "
        "content_sha256 TEXT, byte_size INTEGER);"
        "INSERT INTO jobs VALUES ('job-real');"
    )
    database.close()
    text_source = tmp_path / "appunti.txt"
    markdown_source = tmp_path / "manuale.md"
    text_source.write_text("Contenuto di prova", encoding="utf-8")
    markdown_source.write_text("# Titolo\n\nTesto strutturato", encoding="utf-8")
    request = real_job_request(
        wiki_root,
        "job-real",
        [text_source, markdown_source],
    )
    output_stream = StringIO()

    assert run(StringIO(json.dumps(request) + "\n"), output_stream) == 0
    responses = [json.loads(line) for line in output_stream.getvalue().splitlines()]

    assert responses[-1]["message_type"] == "response"
    assert responses[-1]["payload"]["processed_sources"] == 2
    source_notes = list((wiki_root / "sources").glob("*.md"))
    assert len(source_notes) == 2
    assert all("source_id:" not in note.read_text(encoding="utf-8") for note in source_notes)
    assert not (wiki_root / ".llm-wiki" / "raw").exists()
    with sqlite3.connect(wiki_root / ".llm-wiki" / "catalog.sqlite3") as catalog:
        stored_sources = catalog.execute(
            "SELECT content_sha256, relative_path, path_base FROM source_records "
            "ORDER BY original_name"
        ).fetchall()
    assert len(stored_sources) == 2
    assert all(
        len(row[0]) == 64 and row[2] == text_source.resolve().anchor for row in stored_sources
    )
    assert {row[1] for row in stored_sources} == {
        str(markdown_source.resolve().relative_to(Path(markdown_source.resolve().anchor))),
        str(text_source.resolve().relative_to(Path(text_source.resolve().anchor))),
    }
    assert (wiki_root / ".llm-wiki" / "logs" / "job-real.jsonl").is_file()


def test_worker_extracts_docx_structure(tmp_path: Path) -> None:
    wiki_root = tmp_path / "wiki"
    (wiki_root / ".llm-wiki").mkdir(parents=True)
    (wiki_root / "sources").mkdir()
    database = sqlite3.connect(wiki_root / ".llm-wiki" / "catalog.sqlite3")
    database.executescript(
        "CREATE TABLE jobs (job_id TEXT PRIMARY KEY);"
        "CREATE TABLE source_records ("
        "source_id TEXT PRIMARY KEY, job_id TEXT NOT NULL REFERENCES jobs(job_id), "
        "original_name TEXT NOT NULL, source_format TEXT NOT NULL, "
        "content_sha256 TEXT, byte_size INTEGER);"
        "INSERT INTO jobs VALUES ('job-docx');"
    )
    database.close()
    source = tmp_path / "struttura.docx"
    document = Document()
    document.add_heading("Capitolo", level=1)
    document.add_paragraph("Paragrafo")
    document.save(str(source))
    output_stream = StringIO()

    assert (
        run(
            StringIO(json.dumps(real_job_request(wiki_root, "job-docx", [source])) + "\n"),
            output_stream,
        )
        == 0
    )

    note = next((wiki_root / "sources").glob("*.md")).read_text(encoding="utf-8")
    assert "# Capitolo" in note
    assert "Paragrafo" in note


def test_duplicate_source_updates_one_catalog_record_without_copying_original(
    tmp_path: Path,
) -> None:
    wiki_root = tmp_path / "wiki"
    (wiki_root / ".llm-wiki").mkdir(parents=True)
    (wiki_root / "sources").mkdir()
    database = sqlite3.connect(wiki_root / ".llm-wiki" / "catalog.sqlite3")
    database.executescript(
        "CREATE TABLE jobs (job_id TEXT PRIMARY KEY);"
        "CREATE TABLE source_records ("
        "source_id TEXT PRIMARY KEY, job_id TEXT NOT NULL REFERENCES jobs(job_id), "
        "original_name TEXT NOT NULL, source_format TEXT NOT NULL, "
        "content_sha256 TEXT, byte_size INTEGER);"
        "INSERT INTO jobs VALUES ('job-first');"
        "INSERT INTO jobs VALUES ('job-second');"
    )
    database.close()
    source = tmp_path / "shared.txt"
    source.write_text("Una sola sorgente", encoding="utf-8")

    for job_id in ("job-first", "job-second"):
        output = StringIO()
        assert (
            run(StringIO(json.dumps(real_job_request(wiki_root, job_id, [source])) + "\n"), output)
            == 0
        )

    with sqlite3.connect(wiki_root / ".llm-wiki" / "catalog.sqlite3") as catalog:
        records = catalog.execute(
            "SELECT job_id, relative_path, path_base FROM source_records"
        ).fetchall()
    assert records == [
        (
            "job-second",
            str(source.resolve().relative_to(Path(source.resolve().anchor))),
            source.resolve().anchor,
        )
    ]
    assert not (wiki_root / ".llm-wiki" / "raw").exists()


def test_second_identical_run_reuses_sha256_extraction_cache(tmp_path: Path) -> None:
    wiki_root = tmp_path / "wiki"
    (wiki_root / ".llm-wiki").mkdir(parents=True)
    (wiki_root / "sources").mkdir()
    database = sqlite3.connect(wiki_root / ".llm-wiki" / "catalog.sqlite3")
    database.executescript(
        "CREATE TABLE jobs (job_id TEXT PRIMARY KEY);"
        "CREATE TABLE source_records ("
        "source_id TEXT PRIMARY KEY, job_id TEXT NOT NULL REFERENCES jobs(job_id), "
        "original_name TEXT NOT NULL, source_format TEXT NOT NULL, "
        "content_sha256 TEXT, byte_size INTEGER);"
        "INSERT INTO jobs VALUES ('job-cache-first');"
        "INSERT INTO jobs VALUES ('job-cache-second');"
    )
    database.close()
    source = tmp_path / "cache.md"
    source.write_text("# Contenuto stabile", encoding="utf-8")

    first_output = StringIO()
    assert (
        run(
            StringIO(json.dumps(real_job_request(wiki_root, "job-cache-first", [source])) + "\n"),
            first_output,
        )
        == 0
    )
    artifact = next((wiki_root / ".llm-wiki" / "artifacts").glob("*/document.md"))
    first_modified = artifact.stat().st_mtime_ns

    second_output = StringIO()
    assert (
        run(
            StringIO(json.dumps(real_job_request(wiki_root, "job-cache-second", [source])) + "\n"),
            second_output,
        )
        == 0
    )
    responses = [json.loads(line) for line in second_output.getvalue().splitlines()]

    assert artifact.stat().st_mtime_ns == first_modified
    assert any(response["payload"].get("message") == "source.cache_hit" for response in responses)


def test_worker_persists_a_visible_error_log(tmp_path: Path) -> None:
    wiki_root = tmp_path / "wiki"
    (wiki_root / ".llm-wiki").mkdir(parents=True)
    (wiki_root / "sources").mkdir()
    output_stream = StringIO()
    request = real_job_request(wiki_root, "job-error", [tmp_path / "missing.txt"])

    assert run(StringIO(json.dumps(request) + "\n"), output_stream) == 0
    responses = [json.loads(line) for line in output_stream.getvalue().splitlines()]

    assert responses[-1]["message_type"] == "error"
    assert "detail" in responses[-1]["payload"]
    log_content = (wiki_root / ".llm-wiki" / "logs" / "job-error.jsonl").read_text(encoding="utf-8")
    assert '"level":"error"' in log_content


def test_ocr_monitor_classifies_useful_backend_activity() -> None:
    assert classify_ocr_log_line("Downloading detection model, please wait") == (
        "info",
        "ocr.models_downloading",
        "Downloading detection model, please wait",
    )
    assert classify_ocr_log_line("2026 INFO Accelerator device: 'cpu'") == (
        "info",
        "ocr.accelerator",
        "2026 INFO Accelerator device: 'cpu'",
    )
    assert classify_ocr_log_line("HTTP Request: GET /health") is None
    completed = classify_ocr_log_line("Finished converting document temp.pdf in 12.4 sec.")
    assert completed is not None
    assert completed[1] == "ocr.backend_document_completed"


def test_pdf_ocr_uses_one_full_hybrid_batch_without_page_invocations(tmp_path: Path) -> None:
    sources = [tmp_path / "one.pdf", tmp_path / "two.pdf"]
    command = build_pdf_batch_command(
        tmp_path / "opendataloader-pdf.exe",
        sources,
        tmp_path / "output",
        5002,
    )

    assert command.count(str(sources[0])) == 1
    assert command.count(str(sources[1])) == 1
    assert "--pages" not in command
    assert command.count("--hybrid-mode") == 1
    assert command[command.index("--hybrid-mode") + 1] == "full"


def test_digital_pdfs_use_one_structured_batch_without_ocr(tmp_path: Path) -> None:
    sources = [tmp_path / "one.pdf", tmp_path / "two.pdf"]
    command = build_pdf_direct_batch_command(
        tmp_path / "opendataloader-pdf.exe", sources, tmp_path / "output"
    )

    assert command.count(str(sources[0])) == 1
    assert command.count(str(sources[1])) == 1
    assert "--hybrid" not in command
    assert "--force-ocr" not in command
    assert command[command.index("--format") + 1] == "markdown,json"


def test_ocr_log_reader_waits_for_complete_lines(tmp_path: Path) -> None:
    log_path = tmp_path / "ocr.log"
    log_path.write_bytes(b"Downloading object-detection model")

    offset, entries = read_ocr_log_entries(log_path, 0)
    assert offset == 0
    assert entries == []

    with log_path.open("ab") as stream:
        stream.write(b"\nAccelerator device: cpu\n")
    offset, entries = read_ocr_log_entries(log_path, offset)

    assert offset == log_path.stat().st_size
    assert [entry[1] for entry in entries] == [
        "ocr.models_downloading",
        "ocr.accelerator",
    ]
    assert abs(ocr_progress(5, 10) - 0.8) < 1e-9


def real_job_request(wiki_root: Path, job_id: str, sources: list[Path]) -> dict[str, object]:
    return {
        "protocol_version": "1.0",
        "message_type": "request",
        "request_id": f"request-{job_id}",
        "wiki_id": "wiki-real",
        "job_id": job_id,
        "payload": {
            "action": "start_job",
            "source_paths": [str(source) for source in sources],
            "wiki_root": str(wiki_root),
            "ocr_language": "ita+eng",
        },
    }


def test_transaction_schema_rejects_paths_outside_the_wiki() -> None:
    transaction = read_json("wiki-transaction.json")
    transaction["writes"][0]["relative_path"] = "C:\\Users\\Example\\outside.md"

    errors = list(
        Draft202012Validator(read_schema("wiki-transaction.schema.json")).iter_errors(transaction)
    )

    assert errors


def test_clean_extracted_text_fixes_mojibake_symbols_and_watermarks() -> None:
    raw_corrupted_text = (
        "lOMoARcPSD|68300896\n"
        "L'imprenditore \xc3\xa8 colui che esercita professionalmente un'attivit\xc3\xa0 "
        "economica organizzata "
        "\u00de al fine della produzione o dello scambio di beni o di servizi &gt; art. 2082 c.c.\n"
        "messages.pdf_cover_qr_code_label\n"
        "messages.studocu_not_sponsored_or_endorsed_by_college\n"
        "Downloaded by Mario Rossi (mario@example.com)\n"
        "Studocu non è sponsorizzato o supportato da alcuna università\n"
        "\xc3\x80 \xc3\x88 \xc3\x89 \xc3\x92 \xc3\x99 \xc3\xac \xc3\xb2 \xc3\xb9 \xc3\xa9 "
        "\xe2\x80\x99 \xe2\x80\x9c \xe2\x80\x9d \xe2\x80\x93 \xe2\x80\x94 \xc2\xb0 \xc2\xa7"
    )

    cleaned = clean_extracted_text(raw_corrupted_text)

    assert "lOMoARcPSD" not in cleaned
    assert "messages.pdf_cover" not in cleaned
    assert "messages.studocu" not in cleaned
    assert "Downloaded by" not in cleaned
    assert "Studocu non" not in cleaned
    assert "\xc3\xa8" not in cleaned
    assert "\xc3\xa0" not in cleaned
    assert "\u00de" not in cleaned
    assert "&gt;" not in cleaned
    assert (
        "L'imprenditore è colui che esercita professionalmente un'attività economica organizzata"
        in cleaned
    )
    assert (
        "→ al fine della produzione o dello scambio di beni o di servizi > art. 2082 c.c."
        in cleaned
    )
    assert 'À È É Ò Ù ì ò ù é \' " " \u2013 \u2014 \u00b0 \u00a7' in cleaned


def test_worker_sanitizes_mojibake_during_text_extraction(tmp_path: Path) -> None:
    wiki_root = tmp_path / "wiki"
    (wiki_root / ".llm-wiki").mkdir(parents=True)
    (wiki_root / "sources").mkdir()
    database = sqlite3.connect(wiki_root / ".llm-wiki" / "catalog.sqlite3")
    database.executescript(
        "CREATE TABLE jobs (job_id TEXT PRIMARY KEY);"
        "CREATE TABLE source_records ("
        "source_id TEXT PRIMARY KEY, job_id TEXT NOT NULL REFERENCES jobs(job_id), "
        "original_name TEXT NOT NULL, source_format TEXT NOT NULL, "
        "content_sha256 TEXT, byte_size INTEGER);"
        "INSERT INTO jobs VALUES ('job-clean');"
    )
    database.close()

    source = tmp_path / "appunti-diritto.txt"
    source.write_text(
        "lOMoARcPSD|68300896\n"
        "Nozione di imprenditore: L'imprenditore \xc3\xa8 colui che esercita "
        "un'attivit\xc3\xa0 \u00de commerciale &gt; 2082 c.c.",
        encoding="utf-8",
    )

    output_stream = StringIO()
    assert (
        run(
            StringIO(json.dumps(real_job_request(wiki_root, "job-clean", [source])) + "\n"),
            output_stream,
        )
        == 0
    )

    note = next((wiki_root / "sources").glob("*.md")).read_text(encoding="utf-8")
    assert "lOMoARcPSD" not in note
    assert "\xc3\xa8" not in note
    assert "\xc3\xa0" not in note
    assert "\u00de" not in note
    assert "&gt;" not in note
    assert "L'imprenditore è colui che esercita un'attività → commerciale > 2082 c.c." in note

    artifact = next((wiki_root / ".llm-wiki" / "artifacts").glob("*/document.md")).read_text(
        encoding="utf-8"
    )
    assert "L'imprenditore è colui che esercita un'attività → commerciale > 2082 c.c." in artifact
